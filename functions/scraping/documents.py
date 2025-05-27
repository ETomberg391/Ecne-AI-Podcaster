import os
import PyPDF2 # For PDF processing
import docx # For DOCX processing

from ..utils import log_to_file # Import log_to_file from the utils module

def load_reference_documents(args):
    """
    Loads content from specified reference documents (txt, pdf, docx) or from a folder.
    Returns a list of dictionaries with 'path' and 'content'.
    """
    reference_docs_content = []
    processed_paths = set() # To avoid processing the same file twice if specified by both args

    # --- Load Reference Documents from comma-separated paths ---
    if args.reference_docs:
        print("\nLoading reference documents from paths...")
        log_to_file(f"Attempting to load reference documents from paths: {args.reference_docs}")
        ref_doc_paths = [p.strip() for p in args.reference_docs.split(',') if p.strip()]
        for doc_path in ref_doc_paths:
            full_doc_path = os.path.abspath(doc_path) # Get absolute path for consistent tracking
            if full_doc_path in processed_paths:
                print(f"  - Skipping already processed document: {doc_path}")
                log_to_file(f"Skipping already processed document: {doc_path}")
                continue

            content = None
            try:
                print(f"  - Processing reference document: {doc_path}")
                if doc_path.lower().endswith('.pdf'):
                    # PDF processing
                    text_content = []
                    with open(doc_path, 'rb') as pdf_file: # Open in binary mode
                        reader = PyPDF2.PdfReader(pdf_file) # Use PdfReader
                        if reader.is_encrypted:
                             print(f"    - Warning: Skipping encrypted PDF: {doc_path}")
                             log_to_file(f"Warning: Skipping encrypted PDF: {doc_path}")
                             continue # Skip encrypted PDFs
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text: # Ensure text was extracted
                                text_content.append(page_text)
                    content = "\n".join(text_content)
                    print(f"    - Extracted text from PDF.")
                elif doc_path.lower().endswith('.docx'):
                    # DOCX processing
                    doc = docx.Document(doc_path)
                    text_content = [para.text for para in doc.paragraphs if para.text] # Filter empty paragraphs
                    content = "\n".join(text_content)
                    print(f"    - Extracted text from DOCX.")
                else: # Assume plain text for .txt or unknown/other extensions
                    if not doc_path.lower().endswith('.txt'):
                         print(f"    - Warning: Unknown extension for '{doc_path}', attempting to read as plain text.")
                         log_to_file(f"Warning: Unknown extension for reference doc '{doc_path}', reading as text.")
                    with open(doc_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    print(f"    - Read as plain text.")

                # Process extracted content
                if content and content.strip():
                    reference_docs_content.append({"path": doc_path, "content": content.strip()})
                    processed_paths.add(full_doc_path)
                    print(f"    - Successfully loaded content ({len(content)} chars).")
                    log_to_file(f"Loaded reference doc: {doc_path} ({len(content)} chars)")
                else:
                    print(f"    - Warning: No text content extracted or file is empty: {doc_path}")
                    log_to_file(f"Warning: Reference document {doc_path} empty or no text extracted.")

            except FileNotFoundError:
                print(f"  - Error: Reference document file not found: {doc_path}")
                log_to_file(f"Error: Reference document file not found: {doc_path}")
            except PyPDF2.errors.PdfReadError as pdf_err: # Catch specific PyPDF2 errors
                 print(f"  - Error reading PDF file {doc_path}: {pdf_err}")
                 log_to_file(f"Error reading PDF file {doc_path}: {pdf_err}")
            except Exception as e: # General catch-all
                print(f"  - Error processing reference document {doc_path}: {e}")
                log_to_file(f"Error processing reference document {doc_path}: {e} (Type: {type(e).__name__})")

        if not reference_docs_content and args.reference_docs:
            print("Warning: No valid reference documents were loaded from paths despite --reference-docs being set.")
            log_to_file("Warning: --reference-docs set, but no content loaded from paths.")


    # --- Load Reference Documents from Folder ---
    if args.reference_docs_folder:
        print(f"\nLoading reference documents from folder: {args.reference_docs_folder}")
        log_to_file(f"Attempting to load reference documents from folder: {args.reference_docs_folder}")
        if not os.path.isdir(args.reference_docs_folder):
            print(f"  - Error: Provided path is not a valid directory: {args.reference_docs_folder}")
            log_to_file(f"Error: --reference-docs-folder path is not a directory: {args.reference_docs_folder}")
        else:
            for filename in os.listdir(args.reference_docs_folder):
                doc_path = os.path.join(args.reference_docs_folder, filename)
                full_doc_path = os.path.abspath(doc_path) # Get absolute path for consistent tracking

                if not os.path.isfile(doc_path):
                    continue # Skip subdirectories

                if full_doc_path in processed_paths:
                    print(f"  - Skipping already processed document: {doc_path}")
                    log_to_file(f"Skipping already processed document: {doc_path}")
                    continue

                content = None
                file_ext = os.path.splitext(filename)[1].lower()

                try:
                    print(f"  - Processing reference document: {doc_path}")
                    if file_ext == '.pdf':
                        # PDF processing
                        text_content = []
                        with open(doc_path, 'rb') as pdf_file:
                            reader = PyPDF2.PdfReader(pdf_file)
                            if reader.is_encrypted:
                                print(f"    - Warning: Skipping encrypted PDF: {doc_path}")
                                log_to_file(f"Warning: Skipping encrypted PDF: {doc_path}")
                                continue
                            for page in reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text_content.append(page_text)
                        content = "\n".join(text_content)
                        print(f"    - Extracted text from PDF.")
                    elif file_ext == '.docx':
                        # DOCX processing
                        doc = docx.Document(doc_path)
                        text_content = [para.text for para in doc.paragraphs if para.text]
                        content = "\n".join(text_content)
                        print(f"    - Extracted text from DOCX.")
                    elif file_ext == '.txt':
                        # TXT processing
                        with open(doc_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        print(f"    - Read as plain text.")
                    else:
                        print(f"    - Skipping unsupported file type: {filename}")
                        log_to_file(f"Skipping unsupported file type in reference folder: {filename}")
                        continue # Skip unsupported files

                    # Process extracted content
                    if content and content.strip():
                        reference_docs_content.append({"path": doc_path, "content": content.strip()})
                        processed_paths.add(full_doc_path)
                        print(f"    - Successfully loaded content ({len(content)} chars).")
                        log_to_file(f"Loaded reference doc from folder: {doc_path} ({len(content)} chars)")
                    else:
                        print(f"    - Warning: No text content extracted or file is empty: {doc_path}")
                        log_to_file(f"Warning: Reference document {doc_path} from folder is empty or no text extracted.")

                except FileNotFoundError: # Should not happen with listdir unless race condition
                    print(f"  - Error: Reference document file not found unexpectedly: {doc_path}")
                    log_to_file(f"Error: Reference document file not found unexpectedly: {doc_path}")
                except PyPDF2.errors.PdfReadError as pdf_err:
                    print(f"  - Error reading PDF file {doc_path}: {pdf_err}")
                    log_to_file(f"Error reading PDF file {doc_path} from folder: {pdf_err}")
                except Exception as e:
                    print(f"  - Error processing reference document {doc_path}: {e}")
                    log_to_file(f"Error processing reference document {doc_path} from folder: {e} (Type: {type(e).__name__})")

        log_to_file(f"Finished processing reference documents folder. Total loaded: {len(reference_docs_content)}")

    if not reference_docs_content and (args.reference_docs or args.reference_docs_folder):
         print("Warning: No valid reference documents were loaded from specified paths or folder.")
         log_to_file("Warning: Reference docs/folder specified, but no content loaded.")


    return reference_docs_content